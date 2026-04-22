"""DOCX chapter detection using python-docx."""
from __future__ import annotations
import re
from docx import Document

# Standard Word heading styles
STANDARD_HEADING_STYLES = {"heading 1", "heading 2", "heading1", "heading2", "title"}

# TOC styles to skip
TOC_STYLES = {"toc 1", "toc 2", "toc 3", "toc1", "toc2", "toc3", "contents 1"}

# Regex for lines that are clearly chapter headings
CHAPTER_RE = re.compile(r"^(chapter\s+\d+|chapter\s+[ivxlcdm]+|prologue|epilogue|introduction|conclusion|preface|foreword|afterword)[\s:.\-–—]?",
                        re.IGNORECASE)

# Match tab-number TOC lines like "Chapter 3\t12"
TOC_LINE_RE = re.compile(r".+\t\d+\s*$")


def _is_toc_line(text: str) -> bool:
    return bool(TOC_LINE_RE.match(text))


def _looks_like_chapter_heading(text: str) -> bool:
    t = text.strip()
    if not t or len(t) > 150:
        return False
    return bool(CHAPTER_RE.match(t))


def _detect_chapter_style(doc) -> str | None:
    """
    Find which paragraph style is used for 'Chapter N' lines.
    Returns the style name if found, else None.
    """
    counts: dict[str, int] = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if CHAPTER_RE.match(text) and len(text) < 40:
            style = para.style.name if para.style else ""
            counts[style] = counts.get(style, 0) + 1
    if not counts:
        return None
    # Return the most common style among chapter-heading lines
    return max(counts, key=lambda k: counts[k])


def parse(filepath: str) -> list[dict]:
    doc = Document(filepath)

    # Detect the custom chapter-heading style (e.g. "C1")
    chapter_style = _detect_chapter_style(doc)

    # Also find the "subtitle" style that immediately follows the chapter number
    # (e.g. "C2" in the sample file — the line right after C1)
    subtitle_style = None
    if chapter_style:
        for i, para in enumerate(doc.paragraphs):
            if para.style and para.style.name == chapter_style:
                # Look at the next non-empty paragraph
                for j in range(i + 1, min(i + 4, len(doc.paragraphs))):
                    nxt = doc.paragraphs[j]
                    nxt_text = nxt.text.strip()
                    nxt_style = nxt.style.name if nxt.style else ""
                    if nxt_text and nxt_style != chapter_style and nxt_style.lower() not in TOC_STYLES:
                        subtitle_style = nxt_style
                        break
                if subtitle_style:
                    break

    # Build chapters
    chapters = []
    current_title = None
    current_lines: list[str] = []
    skip_front_matter = True  # skip everything before the first real chapter/intro

    paragraphs = doc.paragraphs
    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        text = para.text.strip()
        style = para.style.name.lower() if para.style else ""

        # Always skip TOC lines
        if style in TOC_STYLES or _is_toc_line(text):
            i += 1
            continue

        # Detect a chapter boundary
        is_chapter_start = False
        chapter_title_parts = []

        if chapter_style and para.style and para.style.name == chapter_style:
            is_chapter_start = True
            chapter_title_parts.append(text)
            # Consume the subtitle line(s) immediately following
            j = i + 1
            while j < len(paragraphs):
                nxt = paragraphs[j]
                nxt_text = nxt.text.strip()
                nxt_style = nxt.style.name if nxt.style else ""
                if not nxt_text:
                    j += 1
                    continue
                if subtitle_style and nxt_style == subtitle_style:
                    chapter_title_parts.append(nxt_text)
                    j += 1
                    break
                break
            i = j
            continue_outer = True
        elif not chapter_style and style in STANDARD_HEADING_STYLES and text:
            is_chapter_start = True
            chapter_title_parts.append(text)
            i += 1
            continue_outer = True
        elif _looks_like_chapter_heading(text):
            # Always catch Introduction/Conclusion/Prologue/etc. regardless of style,
            # because these sections often use a different style than numbered chapters.
            is_chapter_start = True
            chapter_title_parts.append(text)
            i += 1
            continue_outer = True
        else:
            continue_outer = False
            i += 1

        if is_chapter_start:
            # Save previous chapter
            if current_title is not None:
                content = "\n".join(current_lines).strip()
                if content:
                    chapters.append({"title": current_title, "content": content})
            raw_title = ": ".join(chapter_title_parts) if chapter_title_parts else "Chapter"
            current_title = " ".join(raw_title.split())  # collapse whitespace/newlines
            current_lines = []
            skip_front_matter = False
            continue

        if not skip_front_matter and text:
            current_lines.append(text)

    # Flush last chapter
    if current_title and current_lines:
        chapters.append({"title": current_title, "content": "\n".join(current_lines).strip()})

    # Final fallback: return entire doc as one chapter
    if not chapters:
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if full_text:
            chapters.append({"title": "Full Document", "content": full_text})

    return chapters
